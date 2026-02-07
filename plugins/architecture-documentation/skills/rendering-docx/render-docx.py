#!/usr/bin/env python3
"""Architecture docs can be 30-50 pages with 10+ diagrams, and stakeholders
who need to review them often can't work with raw markdown. This script
converts architecture markdown into a polished Word document with all
diagram code blocks rendered as inline images via Kroki.

Parses markdown with mistune's AST mode, renders each diagram code block
through Kroki's REST API (plain-text POST), and builds the DOCX with
python-docx. Detects diagram type from code fence tags and auto-routes
C4 diagrams to the /c4plantuml/ endpoint.
"""

import argparse
import io
import re
import sys
from pathlib import Path

import mistune
import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image


# Every Kroki-supported diagram type that we auto-detect from code fence tags.
# Maps fence tag -> Kroki endpoint path segment.
# Untagged blocks with @startuml are detected separately.
KROKI_DIAGRAM_TYPES = {
    "plantuml", "c4plantuml", "mermaid", "d2", "graphviz", "dot",
    "ditaa", "blockdiag", "seqdiag", "actdiag", "nwdiag", "packetdiag",
    "rackdiag", "bpmn", "bytefield", "excalidraw", "nomnoml", "pikchr",
    "structurizr", "svgbob", "tikz", "umlet", "vega", "vegalite",
    "wavedrom", "wireviz",
}


def parse_args():
    parser = argparse.ArgumentParser(
        description="Convert architecture markdown to DOCX with rendered diagrams"
    )
    parser.add_argument("input", help="Path to input markdown file")
    parser.add_argument("-o", "--output", help="Output DOCX path (default: <input>.docx)")
    parser.add_argument(
        "--kroki-url",
        default="https://kroki.io",
        help="Kroki instance URL (default: https://kroki.io)",
    )
    parser.add_argument(
        "--image-width",
        type=float,
        default=6.0,
        help="Max image width in inches (default: 6.0)",
    )
    parser.add_argument(
        "--template",
        help="Optional .dotx Word template for corporate styling",
    )
    return parser.parse_args()


def render_diagram(code, diagram_type, kroki_url, output_format="png"):
    """Kroki accepts raw diagram source as text/plain POST body and returns
    the rendered image. C4 diagrams need the /c4plantuml/ endpoint, not
    /plantuml/ — Kroki treats them as separate engines.
    """
    # "dot" is an alias for graphviz in Kroki
    endpoint_type = diagram_type
    if endpoint_type == "dot":
        endpoint_type = "graphviz"

    url = f"{kroki_url.rstrip('/')}/{endpoint_type}/{output_format}"
    resp = requests.post(
        url,
        data=code.encode("utf-8"),
        headers={"Content-Type": "text/plain"},
        timeout=30,
    )
    if resp.status_code != 200:
        raise RuntimeError(
            f"Kroki returned {resp.status_code} for {diagram_type}: "
            f"{resp.text[:200]}"
        )
    return resp.content


def detect_diagram_type(fence_tag, code):
    """Determines the Kroki endpoint from the code fence tag and code content.

    Untagged blocks containing @startuml are common in architecture docs.
    C4 diagrams include C4/ headers and need the c4plantuml endpoint —
    sending them to /plantuml/ produces broken output.
    """
    tag = (fence_tag or "").strip().lower()

    if tag in KROKI_DIAGRAM_TYPES:
        return tag

    # Untagged block with @startuml/@enduml — check for C4 markers
    if "@startuml" in code:
        c4_markers = ("!include <C4/", "C4_Context", "C4_Container",
                       "C4_Component", "C4_Deployment")
        if any(marker in code for marker in c4_markers):
            return "c4plantuml"
        return "plantuml"

    return None


def extract_caption(code, preceding_heading, diagram_counter):
    """Caption priority chain designed during brainstorming:
    1. HTML comment <!-- caption: My Caption --> inside the code block
    2. @startuml DiagramName from PlantUML naming
    3. Preceding markdown heading text
    4. Fallback "Diagram N"
    """
    # 1. HTML comment caption
    comment_match = re.search(r'<!--\s*caption:\s*(.+?)\s*-->', code)
    if comment_match:
        return comment_match.group(1)

    # 2. @startuml name
    name_match = re.search(r'@startuml\s+(.+)', code)
    if name_match:
        name = name_match.group(1).strip()
        if name:
            return name

    # 3. Preceding heading
    if preceding_heading:
        return preceding_heading

    # 4. Fallback
    return f"Diagram {diagram_counter}"


def setup_styles(doc):
    """Word's built-in styles don't include figure captions or styled code
    blocks out of the box. We define them once here so the rest of the
    script can just reference style names.
    """
    styles = doc.styles

    # Figure caption style
    if "Figure Caption" not in [s.name for s in styles]:
        caption_style = styles.add_style("Figure Caption", 1)  # 1 = paragraph
        caption_style.font.size = Pt(9)
        caption_style.font.italic = True
        caption_style.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_style.paragraph_format.space_before = Pt(4)
        caption_style.paragraph_format.space_after = Pt(12)

    # Code block style
    if "Code Block" not in [s.name for s in styles]:
        code_style = styles.add_style("Code Block", 1)
        code_style.font.name = "Courier New"
        code_style.font.size = Pt(8.5)
        pf = code_style.paragraph_format
        pf.space_before = Pt(6)
        pf.space_after = Pt(6)
        # Light gray background via shading
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), "F5F5F5")
        code_style.element.get_or_add_pPr().append(shading)


def add_toc(doc):
    """Word TOC is a field code that updates on open. We insert the field
    elements directly because python-docx has no built-in TOC support.
    The user will see "Update this table" on first open in Word.
    """
    paragraph = doc.add_paragraph()
    paragraph.style = doc.styles["Normal"]

    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char_begin)

    run = paragraph.add_run()
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._r.append(instr_text)

    run = paragraph.add_run()
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_char_separate)

    run = paragraph.add_run("Right-click and select Update Field to refresh")
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    run.font.italic = True

    run = paragraph.add_run()
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_end)


def add_image_to_doc(doc, image_bytes, caption, max_width_inches):
    """Constrains diagram images to max_width while maintaining aspect ratio.
    PNG bytes are loaded into a BytesIO stream so python-docx can read
    them without writing temp files to disk.
    """
    stream = io.BytesIO(image_bytes)
    img = Image.open(stream)
    width_px, height_px = img.size
    dpi = img.info.get("dpi", (96, 96))[0] or 96
    width_in = width_px / dpi
    height_in = height_px / dpi

    if width_in > max_width_inches:
        scale = max_width_inches / width_in
        width_in = max_width_inches
        height_in = height_in * scale

    stream.seek(0)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(stream, width=Inches(width_in), height=Inches(height_in))

    cap_para = doc.add_paragraph(caption, style="Figure Caption")
    return cap_para


def add_table_to_doc(doc, header_row, body_rows):
    """Renders markdown tables as Word tables with a header row shaded
    in light blue. Keeps it simple — no merged cells or complex formatting.
    """
    num_cols = len(header_row)
    table = doc.add_table(rows=1, cols=num_cols, style="Table Grid")

    # Header row
    for i, cell_text in enumerate(header_row):
        cell = table.rows[0].cells[i]
        cell.text = cell_text.strip()
        # Bold header text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
        # Light blue shading
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), "D9E2F3")
        cell._tc.get_or_add_tcPr().append(shading)

    # Body rows
    for row_cells in body_rows:
        row = table.add_row()
        for i, cell_text in enumerate(row_cells):
            if i < num_cols:
                row.cells[i].text = cell_text.strip()

    doc.add_paragraph()  # spacing after table


def add_hyperlink(paragraph, url, text):
    """python-docx has no built-in hyperlink support, so we construct
    the w:hyperlink XML element directly. Links need a relationship ID
    pointing to the external URL, plus styled run content.
    """
    part = paragraph.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_style = OxmlElement("w:rStyle")
    r_style.set(qn("w:val"), "Hyperlink")
    r_pr.append(r_style)
    new_run.append(r_pr)

    r_text = OxmlElement("w:t")
    r_text.set(qn("xml:space"), "preserve")
    r_text.text = text
    new_run.append(r_text)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_formatted_paragraph(doc, text, style="Normal"):
    """Handles inline markdown formatting: **bold**, *italic*, `code`,
    and [links](url). Creates a paragraph and adds multiple runs
    for mixed formatting.
    """
    p = doc.add_paragraph(style=style)

    # Pattern: **bold**, *italic*, `code`, [text](url)
    pattern = re.compile(
        r'(\*\*(.+?)\*\*)'   # bold
        r'|(\*(.+?)\*)'       # italic
        r'|(`(.+?)`)'         # inline code
        r'|(\[(.+?)\]\((.+?)\))'  # link
    )

    last_end = 0
    for m in pattern.finditer(text):
        # Add plain text before this match
        if m.start() > last_end:
            p.add_run(text[last_end:m.start()])

        if m.group(2):  # bold
            run = p.add_run(m.group(2))
            run.bold = True
        elif m.group(4):  # italic
            run = p.add_run(m.group(4))
            run.italic = True
        elif m.group(6):  # code
            run = p.add_run(m.group(6))
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        elif m.group(8):  # link
            link_text = m.group(8)
            link_url = m.group(9)
            add_hyperlink(p, link_url, link_text)

        last_end = m.end()

    # Remaining plain text
    if last_end < len(text):
        p.add_run(text[last_end:])

    return p


def process_tokens(doc, tokens, kroki_url, max_width, figure_counter,
                   current_heading):
    """Walks mistune's AST token list and builds the DOCX. Each token type
    maps to a python-docx operation: headings to Heading styles, code
    blocks to either Kroki renders or Code Block style, thematic breaks
    to page breaks, etc.
    """
    for token in tokens:
        tok_type = token.get("type", "")

        if tok_type == "heading":
            level = token.get("attrs", {}).get("level", 1)
            # Extract raw text from children
            text = extract_text_from_children(token.get("children", []))
            current_heading = text
            style = f"Heading {min(level, 4)}"
            doc.add_heading(text, level=min(level, 4))

        elif tok_type == "paragraph":
            text = extract_text_from_children(token.get("children", []))
            if text.strip():
                add_formatted_paragraph(doc, text)

        elif tok_type in ("code_block", "block_code"):
            # mistune v3 may use either token name depending on version
            raw_text = token.get("raw", "") or token.get("text", "")
            attrs = token.get("attrs", {})
            fence_tag = attrs.get("info", "") or ""

            diagram_type = detect_diagram_type(fence_tag, raw_text)

            if diagram_type:
                figure_counter[0] += 1
                caption = extract_caption(
                    raw_text, current_heading, figure_counter[0]
                )
                print(
                    f"  [{figure_counter[0]}] Rendering {diagram_type}: "
                    f"{caption}..."
                )
                image_bytes = render_diagram(
                    raw_text, diagram_type, kroki_url, "png"
                )
                add_image_to_doc(doc, image_bytes, f"Figure {figure_counter[0]}: {caption}", max_width)
            else:
                p = doc.add_paragraph(style="Code Block")
                p.add_run(raw_text)

        elif tok_type == "thematic_break":
            # --- in markdown = page break in DOCX
            doc.add_page_break()

        elif tok_type == "list":
            items = token.get("children", [])
            ordered = token.get("attrs", {}).get("ordered", False)
            for item in items:
                item_text = extract_text_from_children(
                    item.get("children", [])
                )
                # Flatten nested paragraphs in list items
                if item_text.strip():
                    style = "List Number" if ordered else "List Bullet"
                    add_formatted_paragraph(doc, item_text, style=style)

        elif tok_type == "block_quote":
            children = token.get("children", [])
            text = extract_text_from_children_deep(children)
            if text.strip():
                p = doc.add_paragraph(style="Intense Quote")
                p.add_run(text)

        elif tok_type == "table":
            # mistune v3 AST table handling
            process_table_token(doc, token)

        elif tok_type == "blank_line":
            pass  # skip blank lines

    return current_heading


def extract_text_from_children(children):
    """Recursively extracts plain text from mistune AST child nodes."""
    parts = []
    for child in children:
        if isinstance(child, str):
            parts.append(child)
        elif isinstance(child, dict):
            ctype = child.get("type", "")
            if ctype in ("text", "codespan", "raw"):
                raw = child.get("raw", "") or child.get("text", "")
                children_inner = child.get("children", "")
                if raw:
                    parts.append(raw)
                elif children_inner and isinstance(children_inner, str):
                    parts.append(children_inner)
            elif ctype == "paragraph":
                parts.append(
                    extract_text_from_children(child.get("children", []))
                )
            elif ctype == "strong":
                inner = extract_text_from_children(child.get("children", []))
                parts.append(f"**{inner}**")
            elif ctype == "emphasis":
                inner = extract_text_from_children(child.get("children", []))
                parts.append(f"*{inner}*")
            elif ctype == "link":
                link_text = extract_text_from_children(
                    child.get("children", [])
                )
                href = child.get("attrs", {}).get("url", "")
                parts.append(f"[{link_text}]({href})")
            elif ctype == "softbreak":
                parts.append(" ")
            elif ctype == "linebreak":
                parts.append("\n")
            else:
                # Recurse into any other structured node
                inner = child.get("children", [])
                if isinstance(inner, list):
                    parts.append(extract_text_from_children(inner))
    return "".join(parts)


def extract_text_from_children_deep(children):
    """Extracts text from deeply nested structures like block quotes
    which may contain paragraphs with their own children.
    """
    parts = []
    for child in children:
        if isinstance(child, dict):
            ctype = child.get("type", "")
            if ctype == "paragraph":
                parts.append(
                    extract_text_from_children(child.get("children", []))
                )
            elif "children" in child:
                inner = child["children"]
                if isinstance(inner, list):
                    parts.append(extract_text_from_children_deep(inner))
                elif isinstance(inner, str):
                    parts.append(inner)
            elif "raw" in child:
                parts.append(child["raw"])
    return " ".join(parts)


def process_table_token(doc, token):
    """Mistune v3's table plugin produces table_head with table_cell children
    directly (no row wrapper for the header), and table_body with table_row
    children each containing table_cell nodes. We handle both structures.
    """
    children = token.get("children", [])
    header_cells = []
    body_rows = []

    for child in children:
        ctype = child.get("type", "")
        if ctype in ("table_head", "thead"):
            # table_head children may be table_cell directly or table_row
            for item in child.get("children", []):
                itype = item.get("type", "")
                if itype == "table_cell":
                    header_cells.append(
                        extract_text_from_children(item.get("children", []))
                    )
                elif itype in ("table_row", "tr"):
                    for cell in item.get("children", []):
                        header_cells.append(
                            extract_text_from_children(cell.get("children", []))
                        )
        elif ctype in ("table_body", "tbody"):
            for row in child.get("children", []):
                cells = []
                for cell in row.get("children", []):
                    cells.append(
                        extract_text_from_children(cell.get("children", []))
                    )
                body_rows.append(cells)

    if header_cells:
        add_table_to_doc(doc, header_cells, body_rows)


def build_title_page(doc, title, subtitle=None):
    """Creates a clean title page with the document title centered
    vertically. Adds a page break after.
    """
    # Add some spacing to push title toward center
    for _ in range(6):
        doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title)
    run.font.size = Pt(28)
    run.bold = True

    if subtitle:
        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub_para.add_run(subtitle)
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()


def extract_title_from_markdown(text):
    """Pulls the document title from the first H1 heading in the markdown.
    Returns (title, remaining_text) so the H1 isn't duplicated in the body.
    """
    lines = text.split("\n")
    for i, line in enumerate(lines):
        if line.startswith("# ") and not line.startswith("##"):
            title = line[2:].strip()
            remaining = "\n".join(lines[:i] + lines[i + 1:])
            return title, remaining
    return None, text


def main():
    args = parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: {input_path} not found", file=sys.stderr)
        sys.exit(1)

    output_path = args.output or input_path.with_suffix(".docx")
    print(f"render-docx: {input_path} -> {output_path}")
    print(f"Kroki URL: {args.kroki_url}")

    markdown_text = input_path.read_text(encoding="utf-8")

    # Extract title from first H1
    title, markdown_body = extract_title_from_markdown(markdown_text)

    # Create document (with optional template)
    if args.template:
        template_path = Path(args.template)
        if not template_path.exists():
            print(f"Error: template {template_path} not found", file=sys.stderr)
            sys.exit(1)
        doc = Document(str(template_path))
    else:
        doc = Document()

    setup_styles(doc)

    # Title page
    if title:
        build_title_page(doc, title)
    else:
        title = input_path.stem

    # Table of Contents
    doc.add_heading("Table of Contents", level=1)
    add_toc(doc)
    doc.add_page_break()

    # Parse markdown to AST
    md = mistune.create_markdown(renderer="ast", plugins=["table", "strikethrough"])
    tokens = md(markdown_body)

    # Walk tokens and build document
    figure_counter = [0]  # mutable counter for nested function access
    current_heading = None

    print(f"Processing {len(tokens)} top-level blocks...")
    process_tokens(
        doc, tokens, args.kroki_url, args.image_width,
        figure_counter, current_heading
    )

    # Save
    doc.save(str(output_path))
    print(f"\nDone! {figure_counter[0]} diagram(s) rendered.")
    print(f"Output: {output_path}")


if __name__ == "__main__":
    main()
